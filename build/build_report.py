# -*- coding: utf-8 -*-
import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r'C:\Users\Admin03\Desktop\proekt\mirigrushek\report.docx'

doc = Document()

st = doc.styles
st['Normal'].font.name = 'Arial'
st['Normal'].font.size = Pt(11)
for h, sz in [('Heading 1', 16), ('Heading 2', 13.5), ('Heading 3', 11.5)]:
    s = st[h]
    s.font.name = 'Arial'; s.font.size = Pt(sz); s.font.bold = True
    s.font.color.rgb = RGBColor(0, 0, 0)

sec = doc.sections[0]
sec.top_margin = Inches(1); sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1); sec.right_margin = Inches(1)


def set_shading(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def code(text, size=9):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Pt(6)
    set_shading(p, 'F2F2F2')
    lines = text.split('\n')
    for i, line in enumerate(lines):
        r = p.add_run(line)
        r.font.name = 'Consolas'; r.font.size = Pt(size)
        if i < len(lines) - 1:
            r.add_break()
    return p


def para(text, bold=False, italic=False, size=11, align=None, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    return p


def bullets(items):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        p.add_run(it)


def table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].paragraphs[0].add_run(h).bold = True
        set_shading(hdr[i].paragraphs[0], 'F5DEB3')
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].paragraphs[0].add_run(str(val))
    for r in t.rows:
        for i, c in enumerate(r.cells):
            c.width = Inches(widths[i])
    for r in t.rows:
        for c in r.cells:
            for p in c.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def add_page_number(footer_par):
    run = footer_par.add_run()
    fldBegin = OxmlElement('w:fldChar'); fldBegin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    fldEnd = OxmlElement('w:fldChar'); fldEnd.set(qn('w:fldCharType'), 'end')
    run._r.append(fldBegin); run._r.append(instr); run._r.append(fldEnd)


footer = sec.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.add_run('Отчёт о развёртывании ИС «МирИгрушек» — стр. ')
add_page_number(fp)

# ---------------- Титульный лист ----------------
for _ in range(3):
    doc.add_paragraph()
para('Демонстрационный экзамен', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para('по специальности 09.02.07 «Информационные системы и программирование»',
     size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para('Вариант В4 — профильный уровень (вариативная часть)',
     size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
for _ in range(2):
    doc.add_paragraph()
para('ОТЧЁТ О РАЗРАБОТКЕ И РАЗВЁРТЫВАНИИ', bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para('информационной системы магазина игрушек ООО «МирИгрушек»', bold=True, size=16,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
para('Пояснительная записка: использованные команды и соответствие критериям оценки',
     italic=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
for _ in range(6):
    doc.add_paragraph()
para('Дата формирования: ' + datetime.date.today().strftime('%d.%m.%Y'),
     size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ---------------- 1. Общие сведения ----------------
doc.add_heading('1. Общие сведения о работе', level=1)
para('Разработана и развёрнута информационная система магазина игрушек ООО «МирИгрушек» — '
     'веб-приложение на стеке LAMP (Linux, Apache, MySQL, PHP). Система реализует авторизацию '
     'с разграничением прав по ролям, каталог товаров, управление товарами и заказами, регистрацию '
     'новых пользователей. Данные загружаются из предоставленных файлов импорта. Развёртывание на '
     'сервер выполняется одной командой; весь проект размещён в публичном репозитории Git.')
para('Параметры развёрнутой системы:', bold=True, after=2)
table(['Параметр', 'Значение'],
      [['Операционная система', 'Ubuntu Server 22.04 / 24.04 LTS'],
       ['Веб-сервер', 'Apache 2'],
       ['СУБД', 'MySQL 8'],
       ['Язык приложения', 'PHP 8 (модуль libapache2-mod-php)'],
       ['База данных', 'mirigrushek (utf8mb4)'],
       ['Репозиторий', 'https://github.com/shortisshow-prog/mirigrushek'],
       ['Установка одной командой',
        'curl -fsSL https://raw.githubusercontent.com/shortisshow-prog/mirigrushek/main/install.sh | sudo bash'],
       ['Доступ к СУБД (Workbench)', 'host=<IP>, port=3306, user=root, пароль Xmpl123!'],
       ['Вход в приложение', 'логин = фамилия латиницей @mail.ru, пароль Xmpl123!']],
      [2.2, 4.9])

# ---------------- 2. Модуль 1 ----------------
doc.add_heading('2. Модуль 1. Разработка базы данных средствами СУБД', level=1)
para('База данных приведена к третьей нормальной форме (3НФ): выделены справочники, устранены '
     'транзитивные зависимости. Имена объектов заданы по индустриальным стандартам (snake_case). '
     'Состав заказа реализован связью «многие-ко-многим» через таблицу OrderItems.')

para('Перечень таблиц и связей:', bold=True, after=2)
table(['Таблица', 'Назначение', 'Связи (внешние ключи)'],
      [['Roles', 'Роли пользователей', '—'],
       ['Users', 'Пользователи (логин, пароль, ФИО)', 'role_id → Roles'],
       ['Categories', 'Категории товаров', '—'],
       ['Suppliers', 'Поставщики', '—'],
       ['Manufacturers', 'Производители', '—'],
       ['Units', 'Единицы измерения', '—'],
       ['Products', 'Товары', 'unit_id, supplier_id, manufacturer_id, category_id'],
       ['PickupPoints', 'Пункты выдачи', '—'],
       ['OrderStatuses', 'Статусы заказов', '—'],
       ['Orders', 'Заказы', 'pickup_point_id, client_user_id, status_id'],
       ['OrderItems', 'Состав заказа (товар×кол-во)', 'order_id → Orders, product_article → Products']],
      [1.7, 3.2, 4.2])

para('Критерий 1.4 — создание таблиц, полей, первичных ключей и связей. Пример определения '
     'таблицы Products с первичным ключом, типами данных и внешними ключами:', after=2)
code('CREATE TABLE Products (\n'
     '  article         VARCHAR(20) PRIMARY KEY,\n'
     '  name            VARCHAR(500) NOT NULL,\n'
     '  unit_id         INT NOT NULL,\n'
     '  price           DECIMAL(10,2) NOT NULL CHECK (price >= 0),\n'
     '  supplier_id     INT NOT NULL,\n'
     '  manufacturer_id INT NOT NULL,\n'
     '  category_id     INT NOT NULL,\n'
     '  discount        INT NOT NULL DEFAULT 0,\n'
     '  stock_qty       INT NOT NULL DEFAULT 0,\n'
     '  description     TEXT,\n'
     '  photo           VARCHAR(255),\n'
     '  CONSTRAINT fk_prod_cat FOREIGN KEY (category_id) REFERENCES Categories(id)\n'
     ');')

para('Критерий 1.1 — импорт данных. Данные перенесены из файлов импорта (xlsx) в скрипт init.sql '
     'и загружены в БД: 3 роли, 10 пользователей, 4 категории, 5 поставщиков, 3 производителя, '
     '10 товаров, 36 пунктов выдачи, 2 статуса, 10 заказов, 20 позиций заказов. Скрипт init.sql '
     'сформирован автоматически из xlsx генератором generate_sql.py:', after=2)
code('python generate_sql.py        # чтение import/*.xlsx -> init.sql')
para('Критерий 1.2 — скрипт БД. Сформирован единый SQL-скрипт init.sql, создающий БД, все таблицы, '
     'связи и наполняющий их данными. Некорректные данные очищены (например, несуществующая дата '
     '30.02.2025 в заказе №7 загружена как NULL).')
para('Критерий 1.3 — дизайн БД (ERD). ER-диаграмма построена в нотации «вороньи лапки» на основе '
     'приведённой структуры (3НФ, связи 1:N и N:M) и предоставляется отдельным файлом PDF.')

# ---------------- 3. Модуль 2 ----------------
doc.add_heading('3. Модуль 2. Разработка алгоритма и создание приложения', level=1)
bullets([
    '2.5 Авторизация: окно входа — стартовый экран; проверка логина и пароля по таблице Users.',
    '2.6 Отображение ФИО авторизованного пользователя в шапке интерфейса.',
    '2.1–2.4 Приложение реализовано как веб-сайт; идентификаторы и стиль кода — snake_case (PHP).',
    '2.7 Список товаров выводится из БД карточками с фото (или заглушкой picture.png).',
])
para('Критерий 2.8 — подсветка строк (реализовано строго по критериям оценки):', bold=True, after=2)
bullets([
    'при размере скидки более 15% фон карточки товара — цвет #2E8B57;',
    'при сниженной стоимости основная цена перечёркнута и выделена красным, рядом указана '
    'итоговая стоимость чёрным цветом;',
    'если товара нет на складе, строка (карточка) выделяется голубым цветом.',
])
para('Фрагмент логики определения состояния карточки и расчёта итоговой цены (index.php):', after=2)
code("$discount = (int)$p['discount'];\n"
     "$price    = (float)$p['price'];\n"
     "$final    = round($price * (1 - $discount / 100), 2);\n"
     "$outStock = (int)$p['stock_qty'] === 0;\n"
     "$state    = $outStock ? 'oos' : ($discount > 15 ? 'sale' : '');")

# ---------------- 4. Модуль 3 ----------------
doc.add_heading('4. Модуль 3. Последовательный пользовательский интерфейс', level=1)
bullets([
    '3.1–3.5 Формы добавления и редактирования товара: поля для заполнения, загрузка фото '
    '(проверка минимального размера 300×200 px), сохранение в БД, обновление списка.',
    '3.6 Удаление товара с подтверждением; товар, участвующий в заказах, защищён от удаления.',
    '3.7 Поиск по наименованию товара.',
    '3.8 Сортировка по цене (по возрастанию и убыванию).',
    '3.9 Фильтрация по категории.',
    '3.11–3.12 Последовательный интерфейс и сообщения об ошибках при некорректном вводе '
    '(отрицательная цена/количество, неверный формат данных).',
    '3.14 На форме добавления/редактирования присутствуют все поля, указанные в задании.',
])

# ---------------- 5. Модуль 4 ----------------
doc.add_heading('5. Модуль 4. Функционал администратора и менеджера: заказы', level=1)
bullets([
    '4.1 Кнопка «Заказы» открывает список заказов из БД.',
    '4.3 Вывод заказов согласно макету: артикулы заказа, статус, адрес пункта выдачи, '
    'дата заказа, дата доставки, клиент, код получения.',
    '4.4–4.8 Формы добавления и редактирования заказа: статус (выпадающий список), клиент, '
    'пункт выдачи, даты, состав заказа.',
    '4.9 Удаление заказа администратором.',
    '4.10 После операций данные в списке заказов обновляются.',
])

# ---------------- 6. Модуль 5 ----------------
doc.add_heading('6. Модуль 5. Развёртывание серверной инфраструктуры', level=1)
para('Развёртывание полностью автоматизировано. Ниже приведены команды по каждому подкритерию '
     'модуля 5; все они объединены в сценарии install.sh и setup.sh.')

doc.add_heading('6.1 (5.1) Создание и конфигурирование виртуальной машины', level=2)
para('В среде виртуализации (облачная платформа / гипервизор) создаётся виртуальная машина с '
     'ОС Ubuntu Server, сетью в режиме DHCP/мост и доступом по SSH. IP-адрес сервера определяется '
     'командой:')
code('hostname -I')

doc.add_heading('6.2 (5.2) Установка и настройка серверной ОС', level=2)
para('Обновление индексов пакетов и подготовка неинтерактивного режима установки:')
code('export DEBIAN_FRONTEND=noninteractive\napt-get update -y')

doc.add_heading('6.3 (5.3) Установка и настройка СУБД', level=2)
para('Установка MySQL, разрешение сетевых подключений (bind-address = 0.0.0.0), создание БД и '
     'загрузка данных, настройка пользователя root для подключения с любого хоста:')
code('apt-get install -y mysql-server\n'
     'systemctl enable --now mysql\n'
     "sed -i 's/^[[:space:]]*bind-address.*/bind-address = 0.0.0.0/' \\\n"
     '    /etc/mysql/mysql.conf.d/mysqld.cnf\n'
     'systemctl restart mysql\n'
     'mysql -u root --default-character-set=utf8mb4 < init.sql')
para('Настройка root (выполняется в составе init.sql):', after=2)
code("CREATE USER IF NOT EXISTS 'root'@'%' IDENTIFIED BY 'Xmpl123!';\n"
     "ALTER USER 'root'@'%' IDENTIFIED WITH caching_sha2_password BY 'Xmpl123!';\n"
     "GRANT ALL PRIVILEGES ON *.* TO 'root'@'%' WITH GRANT OPTION;\n"
     "FLUSH PRIVILEGES;")

doc.add_heading('6.4 (5.4) Установка и настройка веб-сервера', level=2)
para('Установка Apache и PHP, удаление приветственной страницы Apache и назначение index.php '
     'индексной страницей, чтобы по адресу сервера сразу открывалось приложение:')
code('apt-get install -y apache2 php libapache2-mod-php php-mysql php-mbstring\n'
     'systemctl enable --now apache2\n'
     'rm -f /var/www/html/index.html\n'
     "sed -i 's/DirectoryIndex .*/DirectoryIndex index.php index.html/' \\\n"
     '    /etc/apache2/mods-enabled/dir.conf\n'
     'a2enmod php* \n'
     'systemctl reload apache2')

doc.add_heading('6.5 (5.5) Развёртывание ИС на виртуальной машине', level=2)
para('Получение проекта из репозитория, копирование файлов сайта в каталог веб-сервера, назначение '
     'владельца и открытие портов 80 (сайт) и 3306 (СУБД):')
code('git clone --depth 1 \\\n'
     '    https://github.com/shortisshow-prog/mirigrushek.git /opt/mirigrushek\n'
     'cp -rf /opt/mirigrushek/web/.    /var/www/html/\n'
     'cp -rf /opt/mirigrushek/images/. /var/www/html/images/\n'
     'cp -f  /opt/mirigrushek/report.docx /var/www/html/report.docx\n'
     'chown -R www-data:www-data /var/www/html\n'
     'ufw allow 80/tcp\nufw allow 3306/tcp')
para('Вся последовательность выполняется одной командой на чистом сервере:', after=2)
code('curl -fsSL https://raw.githubusercontent.com/shortisshow-prog/mirigrushek/main/install.sh | sudo bash')

# ---------------- 7. Полные листинги ----------------
doc.add_heading('7. Полные листинги сценариев развёртывания', level=1)
doc.add_heading('7.1 install.sh — загрузчик (клонирование и запуск)', level=2)
code(open(r'C:\Users\Admin03\Desktop\proekt\mirigrushek\install.sh', encoding='utf-8').read(), size=8.5)
doc.add_heading('7.2 setup.sh — установка LAMP, БД и сайта', level=2)
code(open(r'C:\Users\Admin03\Desktop\proekt\mirigrushek\setup.sh', encoding='utf-8').read(), size=8.5)

# ---------------- 8. Сводная таблица команд ----------------
doc.add_heading('8. Сводная таблица использованных команд', level=1)
table(['Команда', 'Назначение'],
      [['python generate_sql.py', 'Генерация init.sql из файлов импорта Excel'],
       ['git init / add / commit', 'Инициализация репозитория и фиксация изменений'],
       ['gh repo create … --push', 'Создание публичного репозитория и публикация на GitHub'],
       ['git push origin main', 'Отправка изменений в удалённый репозиторий'],
       ['curl -fsSL … | sudo bash', 'Установка системы на сервер одной командой'],
       ['apt-get update / install', 'Установка Apache, MySQL, PHP, git'],
       ['systemctl enable --now / restart / reload', 'Управление службами Apache и MySQL'],
       ['sed -i … bind-address', 'Разрешение сетевых подключений к MySQL'],
       ['mysql -u root < init.sql', 'Создание БД и загрузка данных'],
       ['CREATE USER / GRANT', 'Настройка root для подключения отовсюду (Workbench)'],
       ['rm /var/www/html/index.html', 'Удаление приветственной страницы Apache'],
       ['a2enmod php*, DirectoryIndex', 'Назначение index.php индексной страницей'],
       ['cp -rf web/. /var/www/html', 'Развёртывание файлов сайта'],
       ['chown -R www-data', 'Назначение владельца файлов веб-сервера'],
       ['ufw allow 80/3306', 'Открытие портов сайта и СУБД'],
       ['hostname -I', 'Определение IP-адреса сервера']],
      [3.6, 5.5])

# ---------------- 9. Проверка ----------------
doc.add_heading('9. Проверка работоспособности', level=1)
bullets([
    'Сайт открывается по адресу http://<IP-сервера>/ — отображается окно входа, затем каталог.',
    'Подключение в MySQL Workbench: host=<IP>, port=3306, user=root, пароль Xmpl123!.',
    'Вход под ролями (администратор / менеджер / клиент) и проверка прав доступа.',
    'Файл отчёта доступен на сервере: http://<IP-сервера>/report.docx.',
])
para('Ниже размещаются скриншоты корректной работы приложения (вставляются после развёртывания):',
     italic=True, after=2)
for cap in ['Рисунок 1 — окно входа / регистрации',
            'Рисунок 2 — каталог товаров (подсветка скидки и отсутствия на складе)',
            'Рисунок 3 — список заказов',
            'Рисунок 4 — подключение в MySQL Workbench']:
    para('[ место для скриншота ]', align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    para(cap, italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)

doc.save(OUT)
print('saved', OUT)
